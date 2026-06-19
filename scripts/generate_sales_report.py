"""
Generate the FlexxFast EV Monitoring Platform — Sales Proposal Word doc.
Produces: docs/FlexxFast_Platform_Sales_Proposal.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent.parent / "docs" / "FlexxFast_Platform_Sales_Proposal.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ── Default style: Thai-friendly font ───────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Sarabun"
style.font.size = Pt(11)
rpr = style.element.rPr
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
for tag in ("w:eastAsia", "w:cs", "w:hAnsi", "w:ascii"):
    rfonts.set(qn(tag), "Sarabun")

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# Brand colours
BRAND_GREEN  = RGBColor(0x2E, 0x9C, 0x3F)
BRAND_ORANGE = RGBColor(0xF2, 0x8A, 0x1F)
DARK         = RGBColor(0x1F, 0x2A, 0x37)
MUTED        = RGBColor(0x60, 0x6E, 0x80)
ACCENT_BG    = RGBColor(0xF5, 0xF8, 0xFA)

# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def add_title(text, size=26, color=BRAND_GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    return p

def add_h1(text):
    doc.add_paragraph()  # spacing
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = BRAND_GREEN
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = BRAND_ORANGE
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = DARK
    return p

def add_para(text, size=11, color=None, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def add_bullets(items, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            head, body = it
            run = p.add_run(head)
            run.bold = True
            run.font.size = Pt(size)
            run2 = p.add_run("  " + body)
            run2.font.size = Pt(size)
        else:
            run = p.add_run(it)
            run.font.size = Pt(size)

def add_feature_box(title, lines):
    """Single-cell highlighted box for a feature description."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "F5F8FA")
    # Title
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(4)
    r0 = p0.add_run(title)
    r0.font.size = Pt(12)
    r0.font.bold = True
    r0.font.color.rgb = BRAND_GREEN
    for ln in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("• " + ln)
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
    doc.add_paragraph()

def add_page_section_table(title, what_user_sees, business_value):
    tbl = doc.add_table(rows=3, cols=2)
    tbl.autofit = True
    # Row 0 — title spanning both
    title_cell = tbl.rows[0].cells[0]
    title_cell.merge(tbl.rows[0].cells[1])
    set_cell_bg(title_cell, "2E9C3F")
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Row 1 — "ผู้ใช้เห็นอะไร"
    c10 = tbl.rows[1].cells[0]
    c11 = tbl.rows[1].cells[1]
    set_cell_bg(c10, "F5F8FA")
    p = c10.paragraphs[0]
    r = p.add_run("สิ่งที่ผู้ใช้งานเห็น")
    r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = MUTED
    for line in what_user_sees:
        p = c11.add_paragraph() if c11.paragraphs[0].text else c11.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("• " + line)
        r.font.size = Pt(11)
    # Row 2 — "คุณค่าทางธุรกิจ"
    c20 = tbl.rows[2].cells[0]
    c21 = tbl.rows[2].cells[1]
    set_cell_bg(c20, "F5F8FA")
    p = c20.paragraphs[0]
    r = p.add_run("คุณค่าทางธุรกิจ")
    r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = MUTED
    for line in business_value:
        p = c21.add_paragraph() if c21.paragraphs[0].text else c21.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("• " + line)
        r.font.size = Pt(11)
    # Column widths
    for row in tbl.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(12.5)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FlexxFast")
r.font.size = Pt(54)
r.font.bold = True
r.font.color.rgb = BRAND_GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("EV Charger Monitoring Platform")
r.font.size = Pt(20)
r.font.color.rgb = BRAND_ORANGE
r.font.bold = True

doc.add_paragraph()
add_title("แพลตฟอร์มมอนิเตอร์สถานีชาร์จ EV แบบเรียลไทม์", size=15, color=DARK, bold=False, space_after=2)
add_title("เห็นทุกสถานี รู้ทุกปัญหา ก่อนลูกค้าจะรู้", size=13, color=MUTED, bold=False)

for _ in range(8):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sales Proposal Document")
r.font.size = Pt(13)
r.font.bold = True
r.font.color.rgb = DARK

add_title("by EDS Engineering Solutions", size=12, color=MUTED, bold=False)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 1. บทนำ
# ════════════════════════════════════════════════════════════════════════

add_h1("1. บทนำ — FlexxFast คืออะไร")

add_para(
    "FlexxFast คือแพลตฟอร์มมอนิเตอร์สถานีชาร์จ EV แบบเว็บแอปพลิเคชัน (Web Dashboard) "
    "ออกแบบมาสำหรับผู้ดำเนินการสถานีชาร์จที่ต้องการเห็นสถานะทุกตู้ในมือเดียว ตั้งแต่ตู้เดียว "
    "ไปจนถึงเครือข่ายหลักร้อยตู้ ระบบรับข้อมูลแบบเรียลไทม์ผ่าน MQTT จากตัวควบคุมในตู้ "
    "(รองรับทั้ง Phoenix Controller รุ่นปัจจุบัน และ Vector Controller รุ่นใหม่) แล้วประมวลผล "
    "เก็บประวัติใน MongoDB ส่งต่อให้แดชบอร์ดแสดงผลสวยงามใช้งานง่าย"
)

add_para(
    "เป้าหมายของแพลตฟอร์ม: ลดเวลาตรวจหาความผิดปกติของตู้ ลดเวลา downtime ของบริการ และ "
    "ทำให้ทีมงานสามารถตรวจสอบสถานะตู้ได้ตลอดเวลา จากที่ทำงาน ที่บ้าน หรือบนโทรศัพท์มือถือ"
)

add_h2("FlexxFast แก้ปัญหาอะไร?")
add_bullets([
    ("ไม่รู้ว่าตู้เสีย ก่อนลูกค้าจะโทรมา", "ระบบแจ้งเตือนทันทีเมื่อตู้ไม่ส่งสัญญาณ"),
    ("ไม่รู้ว่าตู้ตัวไหน เครื่องใดมีปัญหา", "หน้าจอแสดงสถานะทุกตู้พร้อมกัน เห็นตู้แดงทันที"),
    ("ตรวจสอบยาก ต้องไปดูที่หน้าตู้", "เข้าหน้าเว็บที่ไหนก็ได้ ทั้งคอมพิวเตอร์และมือถือ"),
    ("ไม่มีข้อมูลการใช้งานย้อนหลัง", "ระบบเก็บประวัติทุกรายการ ดูยอดการชาร์จและพลังงานย้อนหลังได้"),
    ("ไม่รู้ว่าตู้ทำงานเต็มประสิทธิภาพหรือเปล่า", "ดูแรงดัน กระแส กำลัง อุณหภูมิ ครบทุกค่า"),
])

# ════════════════════════════════════════════════════════════════════════
# 2. คุณสมบัติเด่น
# ════════════════════════════════════════════════════════════════════════

add_h1("2. คุณสมบัติเด่นของระบบ (Key Features)")

add_feature_box("เรียลไทม์ ไม่ต้องรีเฟรช", [
    "ข้อมูลอัปเดตทันทีผ่าน WebSocket — ไม่ต้องกด F5",
    "เห็นการเปลี่ยนสถานะของตู้ทุกตู้พร้อมกัน",
    "ความล่าช้าวัดเป็นวินาที ไม่ใช่นาที",
])

add_feature_box("รองรับหลายรุ่นตู้ในระบบเดียว", [
    "Phoenix Controller (รุ่นปัจจุบัน) — ใช้ topic แยกตามอุปกรณ์",
    "Vector Controller (รุ่นใหม่) — ใช้ state topic เดียวรวมทุกค่า",
    "เปลี่ยนรุ่นได้ในหน้า Config ไม่ต้องเปลี่ยนระบบ",
])

add_feature_box("ระบบแจ้งเตือนความปลอดภัย", [
    "อุณหภูมิสายชาร์จเกินขีดอันตราย (≥ 190 °C)",
    "ตรวจจับการกดปุ่มฉุกเฉิน (Emergency Stop)",
    "Power Module หลุดออกจากระบบ",
    "ไม่มี Heartbeat เกิน 5 นาที = ตู้ออฟไลน์",
    "Meter ค้าง > 2 วัน = ไม่มีการใช้งานจริง",
])

add_feature_box("รายงานพลังงานครบครัน", [
    "นับจำนวนการชาร์จรายวัน รายเดือน",
    "พลังงานรวม (kWh) ที่ขายได้",
    "แยกตามตู้ แยกตามหัวชาร์จ",
    "Export ข้อมูลย้อนหลังได้ 30 วัน",
])

add_feature_box("ใช้งานง่าย รองรับมือถือ", [
    "Sidebar เมนูตามหมวด ไม่ต้องค้นหา",
    "ค้นหาสถานีด้วยชื่อ — auto-complete",
    "เรียงตู้ตามชื่อ (A-Z) หรือเรียง 'ตู้ที่มีปัญหาก่อน'",
    "รองรับ Dark Mode / Light Mode ปรับสีตามใจชอบ",
    "เปิดบนมือถือได้ทันที (Responsive Design)",
])

add_feature_box("ความปลอดภัยและสิทธิ์การเข้าถึง", [
    "ระบบ Login (Username + Password) ก่อนเข้าทุกหน้า",
    "Session มีอายุ 7 วัน รักษาความปลอดภัยด้วยลายเซ็น HMAC",
    "หน้า Config (เพิ่ม/ลบตู้) มี PIN gate ป้องกันการแก้ไขโดยไม่ได้รับอนุญาต",
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 3. ภาพรวมโครงสร้างหน้าเว็บ
# ════════════════════════════════════════════════════════════════════════

add_h1("3. ภาพรวมโครงสร้างหน้าเว็บ")

add_para(
    "ระบบจัดเมนูเป็น 3 หมวดหลัก แสดงในแถบด้านซ้าย (Sidebar) ตามลำดับความถี่ในการใช้งาน:"
)

# Menu structure table
tbl = doc.add_table(rows=4, cols=2)
tbl.autofit = True
headers = [
    ("หมวด Navigation", ["Overview (ภาพรวมทุกตู้)", "Alert Center (ศูนย์การแจ้งเตือน)", "Station Config (เพิ่ม/แก้ไขตู้)", "Settings (ตั้งค่าทั่วไป)"]),
    ("หมวด Statistics", ["Energy Usage (รายงานพลังงาน + จำนวนการชาร์จ)"]),
    ("หมวด System Overview", ["Heartbeat (ตรวจสถานะอุปกรณ์ในตู้)", "Power Module (จำนวนโมดูลที่ใช้งาน)", "Meter (มิเตอร์พลังงาน)", "Temperature (อุณหภูมิ Router)", "Fan RPM (รอบพัดลม)", "Device Status (สถานะอุปกรณ์ + อุณหภูมิสาย)", "MQTT Scripts (สถานะสคริปต์เบื้องหลัง)"]),
    ("หมวด Station Detail", ["รายละเอียดในตู้ — 7 แท็บข้อมูลของตู้นั้นๆ"]),
]
for i, (cat, items) in enumerate(headers):
    c0 = tbl.rows[i].cells[0]
    c1 = tbl.rows[i].cells[1]
    set_cell_bg(c0, "F5F8FA")
    p = c0.paragraphs[0]
    r = p.add_run(cat)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND_GREEN
    for j, item in enumerate(items):
        p = c1.add_paragraph() if j > 0 else c1.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("• " + item)
        r.font.size = Pt(10)
    tbl.rows[i].cells[0].width = Cm(5.0)
    tbl.rows[i].cells[1].width = Cm(11.5)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════
# 4. หน้าจอแต่ละหน้า
# ════════════════════════════════════════════════════════════════════════

add_h1("4. หน้าจอแต่ละหน้า — ทำอะไรได้บ้าง")

# ─── Login ───
add_h2("4.1 หน้า Login")
add_page_section_table(
    "ประตูทางเข้าระบบ (Sign in)",
    [
        "ช่องกรอก Username + Password",
        "โลโก้ FlexxFast แสดงตรงกลาง",
        "ข้อความแจ้งเตือนชัดเจนเมื่อกรอกผิด",
    ],
    [
        "ป้องกันบุคคลภายนอกเข้าถึงข้อมูลตู้ของคุณ",
        "Session อายุ 7 วัน — ไม่ต้อง login ทุกครั้ง",
        "รองรับการแยก user หลายคน (สามารถขยายได้)",
    ],
)

# ─── Overview (Home) ───
add_h2("4.2 หน้า Overview (หน้าแรก)")
add_page_section_table(
    "ภาพรวมทุกตู้ในเครือข่าย",
    [
        "การ์ดสรุปแต่ละสถานี — ชื่อตู้ สถานะ (Online / Degraded / Offline)",
        "สีขอบการ์ด: เขียว = ปกติ, เหลือง = ผิดปกติบางส่วน, แดง = ตู้ดับ",
        "ตัวเลขสรุปด้านบน: Total / Online / Degraded / Offline",
        "ปุ่ม Search หาตู้ตามชื่อได้ทันที",
        "ปุ่ม A-Z หรือ Problems First (จัดตู้ที่มีปัญหาขึ้นบน)",
    ],
    [
        "เห็นภาพรวมเครือข่ายในจอเดียว ไม่ต้องคลิกเข้าทุกตู้",
        "หาตู้ที่มีปัญหาได้ในวินาทีเดียว",
        "เหมาะกับการเปิดทิ้งไว้บนจอใหญ่ในห้องควบคุม",
    ],
)

# ─── System Overview / Heartbeat ───
add_h2("4.3 หน้า Heartbeat")
add_page_section_table(
    "สถานะการสื่อสารของอุปกรณ์ในตู้",
    [
        "แสดงสถานะของ 3 อุปกรณ์ในแต่ละตู้: OCPP Device, Pi5, Router",
        "บอกเวลาที่เห็นล่าสุด (last seen) แต่ละอุปกรณ์",
        "สรุปด้านบน: X/Y Online สำหรับแต่ละประเภท",
        "สีตัวเลขเปลี่ยนตามเปอร์เซ็นต์: ≥90% เขียว, ≥80% เหลือง, <80% แดง",
    ],
    [
        "ดูได้ทันทีว่าอุปกรณ์ตัวไหนของตู้ไหนหลุดออกจากระบบ",
        "ช่วยทีมเทคนิคพุ่งตรงไปที่ตู้ที่มีปัญหาได้เลย",
    ],
)

# ─── Power Module ───
add_h2("4.4 หน้า Power Module")
add_page_section_table(
    "จำนวนโมดูลกำลังที่ใช้งานต่อหัวชาร์จ",
    [
        "แสดงจำนวนโมดูลที่ออนไลน์ vs. ที่ตั้งคาดหวังไว้",
        "Badge: FULL (ครบ) / INCOMPLETE (ขาดบางตัว)",
        "สถานะการชาร์จต่อหัว (Ready / Connected / Charging / Fault)",
        "แรงดันและกระแส realtime ของแต่ละกลุ่ม",
    ],
    [
        "รู้ทันทีว่าตู้ไหนกำลังจ่ายไฟไม่เต็มที่เพราะโมดูลตัวใดตัวหนึ่งหลุด",
        "วางแผนซ่อมบำรุงโมดูลล่วงหน้าได้",
    ],
)

# ─── Meter ───
add_h2("4.5 หน้า Meter")
add_page_section_table(
    "มิเตอร์พลังงานสะสมของแต่ละหัวชาร์จ",
    [
        "ตัวเลข kWh สะสมต่อหัว ทั้ง 2 หัว",
        "LED สถานะ: เขียว = มีการใช้งาน, แดง = ค่าค้างเกิน 2 วัน",
        "Timestamp ของการเปลี่ยนค่าครั้งล่าสุด (เก็บข้ามรีสตาร์ทระบบ)",
        "ตู้ที่ค่าค้าง = อาจมีปัญหา meter หรือลูกค้าไม่ใช้บริการ",
    ],
    [
        "ตรวจสอบรายได้รายตู้ได้ทันที",
        "พบตู้ที่ไม่มียอดขายเลยใน 2 วัน → ส่งคนไปตรวจสอบ",
    ],
)

# ─── Temperature ───
add_h2("4.6 หน้า Temperature")
add_page_section_table(
    "อุณหภูมิ Router ในตู้",
    [
        "ค่า °C ของ Router ในแต่ละตู้",
        "เกจสีไล่ระดับ: เขียว < 70°C, เหลือง 70-79°C, แดง ≥ 80°C",
        "สรุป Avg / Max / Alerts ในสรุปด้านบน",
        "ถ้าตู้ไม่ส่งข้อมูล > 5 นาที จะแสดง 'NO DATA'",
    ],
    [
        "ตรวจจับตู้ที่อุณหภูมิภายในสูงผิดปกติ — ป้องกัน Router เสียก่อนถึงเวลา",
        "ลดความเสี่ยงตู้ดับเพราะอุปกรณ์ภายในร้อนเกิน",
    ],
)

# ─── Fan RPM ───
add_h2("4.7 หน้า Fan RPM")
add_page_section_table(
    "รอบการทำงานของพัดลมระบายความร้อน",
    [
        "FAN 1 ถึง FAN 8 ของแต่ละตู้ พร้อมตัวเลข RPM",
        "ระบุยี่ห้อพัดลมที่ใช้ (EBM / Winstrom / DAKO)",
        "แสดงอุณหภูมิ Router ในมุมขวาบนของแต่ละการ์ด",
        "พัดลมที่หยุด (0 RPM) จะเห็นเด่นชัด",
    ],
    [
        "ตรวจจับพัดลมเสียได้แม้ตู้ยังทำงานอยู่",
        "วางแผนเปลี่ยนพัดลมก่อนระบบ Overheat",
    ],
)

# ─── Device Status (Vector + safety) ───
add_h2("4.8 หน้า Device Status (สำหรับตู้ Vector)")
add_page_section_table(
    "ภาพรวมความปลอดภัยของตู้ — สถานะอุปกรณ์ + อุณหภูมิสายชาร์จ",
    [
        "Status pill: HMI / PLC1 / PLC2 / IMD (ตัวตรวจฉนวน)",
        "อุณหภูมิขั้วบวก/ลบ ของหัวชาร์จทั้ง 2 หัว",
        "ป้าย DANGER ทันทีหาก: อุปกรณ์ Inactive / อุณหภูมิ ≥ 190°C / กดปุ่มฉุกเฉิน",
        "ป้าย WARNING เมื่อ: อุณหภูมิ 150-189°C / มี temperature sensor เสีย",
        "แสดง 'EMERGENCY STOP ACTIVE' ด้วยแถบสีแดงเด่นชัด",
    ],
    [
        "เห็นภาพรวมความปลอดภัยทุกตู้ในจอเดียว",
        "ตรวจจับตู้ที่เสี่ยงไฟไหม้สายชาร์จก่อนเกิดเหตุ",
        "Compliance — เก็บประวัติเหตุการณ์ฉุกเฉินครบถ้วน",
    ],
)

# ─── MQTT Scripts ───
add_h2("4.9 หน้า MQTT Scripts")
add_page_section_table(
    "สถานะของสคริปต์เบื้องหลังในตู้",
    [
        "ดูว่า fault_status / plc script ยังทำงานปกติหรือไม่",
        "เวลา Heartbeat ล่าสุดของแต่ละสคริปต์",
        "Badge ALL RUNNING / STOPPED",
    ],
    [
        "ตรวจจับปัญหาในซอฟต์แวร์ฝั่ง edge ได้ก่อนที่ตู้จะหยุดทำงาน",
        "เหมาะสำหรับทีมเทคนิคในการ debug",
    ],
)

# ─── Statistics — Energy Usage ───
add_h2("4.10 หน้า Energy Usage (รายงานการขายพลังงาน)")
add_page_section_table(
    "รายงานพลังงานและจำนวนการชาร์จย้อนหลัง",
    [
        "เลือกสถานี (ค้นหาในรายการได้)",
        "เลือกช่วงเวลา: 7 วัน / 30 วัน / เดือนนี้",
        "Summary: Total Sessions / Total kWh / Avg per Day",
        "กราฟ Line chart แสดง kWh รายวัน (Head 1 + Head 2)",
        "กราฟ Line chart แสดง Sessions รายวัน",
        "ตารางรายละเอียดรายวัน — H1/H2 sessions + kWh + total",
        "Tooltip อ่านค่าตามตำแหน่งเมาส์",
    ],
    [
        "นำเสนอรายได้ให้ผู้บริหารได้ทันที",
        "วิเคราะห์ตู้ที่ขายได้น้อย — พิจารณาย้ายตำแหน่ง",
        "ใช้วางแผนการลงทุนตู้เพิ่มในจุดที่มียอดสูง",
    ],
)

# ─── Station Detail ───
add_h2("4.11 หน้า Station Detail — รายตู้")
add_page_section_table(
    "เข้าดูข้อมูลในตู้ตู้เดียวอย่างละเอียด (มี 7 แท็บ)",
    [
        "Heartbeat — สถานะอุปกรณ์ในตู้",
        "Power Module — โมดูลกำลังแยกหัวชาร์จ",
        "Meter — มิเตอร์ + กราฟ 48 ชม.",
        "Temperature — อุณหภูมิ Router + กราฟ 24 ชม.",
        "Fan RPM — รอบพัดลม 8 ตัว + เกจวงกลม",
        "MQTT Scripts — script monitoring",
        "PLC — รายละเอียดการชาร์จ (ระดับ low-level)",
    ],
    [
        "Drill-down จากภาพรวมเข้าไปดูตู้ใดตู้หนึ่งโดยเฉพาะ",
        "ใช้ตอน Troubleshoot — ดู context ทั้งหมดของตู้ตัวนั้นในจอเดียว",
    ],
)

# ─── Alert Center ───
add_h2("4.12 หน้า Alert Center")
add_page_section_table(
    "รวมการแจ้งเตือนทั้งหมดในที่เดียว",
    [
        "Filter: ดู Unacknowledged เท่านั้น / ดูทั้งหมด",
        "เรียงตามเวลา ใหม่สุดอยู่บน",
        "แต่ละ alert: เวลา / สถานี / ระดับความรุนแรง (warning, critical) / ข้อความ",
        "ปุ่ม Acknowledge — กดเพื่อจัดการแล้ว",
    ],
    [
        "ทีมงานเห็นว่ามีอะไรเกิดขึ้นกับตู้บ้าง โดยไม่ต้องเปิดทุกหน้า",
        "Audit trail — ใครจัดการ alert ไหน เมื่อไหร่",
    ],
)

# ─── Station Config ───
add_h2("4.13 หน้า Station Config")
add_page_section_table(
    "หน้าเพิ่ม / แก้ไข / ลบสถานี (มี PIN gate)",
    [
        "รายการตู้ทั้งหมด — Search ได้, แสดง topic หลัก, brand",
        "ปุ่ม + Add Station เพิ่มตู้ใหม่",
        "ปุ่ม Edit แก้ไข Topic, Collection, Telegram, Brand, Controller Type",
        "ปุ่ม Delete (ต้องยืนยันก่อน) ลบตู้",
        "Validate ชื่อ (alphanumeric เท่านั้น) ป้องกัน MongoDB เสียหาย",
        "ตรวจซ้ำชื่อ — ห้ามใส่ชื่อตู้ซ้ำ",
        "Auto-cleanup: เปลี่ยนชื่อตู้ → ลบ collection เก่าให้",
    ],
    [
        "ทีมเทคนิคจัดการการเพิ่ม/ลดตู้ได้เอง ไม่ต้องเข้า Mongo",
        "ขยายธุรกิจได้คล่อง — เปิดตู้ใหม่ปุ๊บใช้ปั๊บ",
    ],
)

# ─── Settings ───
add_h2("4.14 หน้า Settings")
add_page_section_table(
    "ตั้งค่าระบบรวม",
    [
        "ค่าขีดอันตรายของอุณหภูมิ",
        "ตั้งค่า Telegram bot สำหรับแจ้งเตือน",
        "ค่าทั่วไปอื่นๆ",
    ],
    [
        "ปรับให้เหมาะกับแต่ละพื้นที่ (เช่น สาขาในที่อากาศร้อน อาจตั้ง threshold สูงกว่า)",
    ],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 5. ระบบแจ้งเตือน
# ════════════════════════════════════════════════════════════════════════

add_h1("5. ระบบแจ้งเตือนอัตโนมัติ")

add_para(
    "FlexxFast มี Auto-Alert Engine ทำงานเบื้องหลังตลอด 24 ชั่วโมง ตรวจสอบเงื่อนไขความผิดปกติ "
    "และสร้างการแจ้งเตือนทันทีเมื่อพบ:"
)

add_h3("ประเภทการแจ้งเตือน")
add_bullets([
    ("🌡️ Temperature Alert", "อุณหภูมิ Router เกิน 80°C → critical"),
    ("💓 Heartbeat Offline", "อุปกรณ์ไม่ส่งสัญญาณเกิน 5 นาที → critical"),
    ("⚡ Power Module Missing", "จำนวนโมดูลออนไลน์น้อยกว่าที่ตั้งคาด → warning"),
    ("📊 Meter Stalled", "มิเตอร์ค่าไม่เปลี่ยนเกิน 2 วัน → warning"),
    ("🔧 Script Offline", "สคริปต์ใน edge หยุดทำงานเกิน 5 นาที → warning"),
    ("🚨 Emergency Stop", "ตู้ Vector — กดปุ่มฉุกเฉิน → critical"),
])

add_h3("คุณสมบัติ Alert Engine")
add_bullets([
    "Dedupe: alert เดียวกันไม่ยิงซ้ำ ตราบใดที่เงื่อนไขยังเป็นจริง",
    "Clear: เมื่อเงื่อนไขหายไป alert ปิดอัตโนมัติ — alert ถัดไปยิงใหม่ได้",
    "Persist: state ของ alert เก็บใน MongoDB — รอดข้าม restart",
    "Telegram Integration: ส่งข้อความเข้ากลุ่ม Telegram ของลูกค้าได้",
    "Acknowledge: ผู้ดูแลกดยอมรับ alert แล้วลด noise ในหน้าจอ",
])

# ════════════════════════════════════════════════════════════════════════
# 6. ความปลอดภัย
# ════════════════════════════════════════════════════════════════════════

add_h1("6. ความปลอดภัยและการเข้าถึง")

add_h3("ระบบล็อกอิน 2 ระดับ")
add_bullets([
    ("ระดับ 1 — Login เข้าใช้งาน", "Username + Password ก่อนเข้าทุกหน้า ใช้ HMAC-signed cookie ปลอมไม่ได้"),
    ("ระดับ 2 — PIN gate หน้า Config", "สำหรับแอดมิน — ป้องกันการแก้ไขโดยไม่ได้รับอนุญาต"),
])

add_h3("คุณสมบัติด้านความปลอดภัย")
add_bullets([
    "Session อายุ 7 วัน — ปลอดภัยพอ ใช้งานสะดวก",
    "HttpOnly cookie — Browser script เข้าถึงไม่ได้ ลดความเสี่ยง XSS",
    "API ทุกตัวกั้นด้วย middleware — แม้เปิด URL ตรงๆ ก็โดน 401",
    "พร้อมรองรับ HTTPS / Reverse Proxy (nginx, Caddy)",
])

# ════════════════════════════════════════════════════════════════════════
# 7. โครงสร้างเชิงเทคนิค
# ════════════════════════════════════════════════════════════════════════

add_h1("7. สถาปัตยกรรมเทคนิค (ภาพรวม)")

add_para(
    "FlexxFast ใช้สถาปัตยกรรม 3-layer ที่ทนทาน scale ได้ และดูแลรักษาง่าย:"
)

tbl = doc.add_table(rows=3, cols=2)
tbl.autofit = True
layers = [
    ("🟢 Edge Layer", "MQTT message broker — รับข้อมูลจากตู้ทุกตัวแบบ pub/sub. Backend service (Node.js) subscribe → process → ทำการแจ้งเตือน → broadcast WS ให้ Frontend"),
    ("🟢 Storage Layer", "MongoDB — เก็บ config ของตู้, ข้อมูลย้อนหลัง, log การแจ้งเตือน, รายงานพลังงาน. รองรับ replication และ backup"),
    ("🟢 Presentation Layer", "Next.js Web App — แสดงผลแบบ Server-Side Rendering, Real-time WebSocket bridge, รองรับการเปิดบนมือถือ"),
]
for i, (name, desc) in enumerate(layers):
    c0 = tbl.rows[i].cells[0]
    c1 = tbl.rows[i].cells[1]
    set_cell_bg(c0, "F5F8FA")
    p = c0.paragraphs[0]
    r = p.add_run(name)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND_GREEN
    p = c1.paragraphs[0]
    r = p.add_run(desc)
    r.font.size = Pt(11)
    tbl.rows[i].cells[0].width = Cm(5.0)
    tbl.rows[i].cells[1].width = Cm(11.5)

doc.add_paragraph()
add_h3("เทคโนโลยีที่ใช้")
add_bullets([
    "Next.js 16 + TypeScript + React 19 (Frontend)",
    "Node.js + tsx + MQTT.js + WebSocket (Backend)",
    "MongoDB (Database)",
    "Web Crypto API (Authentication)",
])

add_h3("ความสามารถในการขยาย")
add_bullets([
    "รองรับสถานีได้มากกว่า 200 ตู้ในระบบเดียว (ทดสอบจริงแล้ว)",
    "Auto-restart resilience — ระบบฟื้นตัวเองเมื่อมีปัญหา",
    "Backpressure protection — รับมือกับ MQTT burst ได้",
    "On-demand backfill — เพิ่มตู้ใหม่ → ข้อมูลย้อนหลัง 30 วันโหลดอัตโนมัติ",
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 8. ทำไมต้อง FlexxFast
# ════════════════════════════════════════════════════════════════════════

add_h1("8. ทำไมต้องเลือก FlexxFast")

reasons = [
    ("ออกแบบโดยทีมที่เข้าใจ EV charger จริงๆ",
     "เราพัฒนาตู้ EV และระบบมอนิเตอร์ควบคู่กันมาตลอด ไม่ใช่ระบบทั่วไปที่ดัดแปลงมา"),
    ("รองรับ Controller หลายรุ่น",
     "Phoenix รุ่นปัจจุบัน + Vector รุ่นใหม่ — เพิ่มได้อีกในอนาคต"),
    ("ใช้งานง่ายจริง",
     "UI ออกแบบให้พนักงานสามารถใช้ได้ทันที ไม่ต้องเทรนนิ่งหลายวัน"),
    ("เห็นปัญหาก่อนลูกค้า",
     "Auto-alert + Telegram → ทีมตอบสนองเร็ว ลด downtime"),
    ("รายงานพลังงานพร้อมนำเสนอ",
     "Export กราฟและตัวเลขให้ผู้บริหารได้ทันที"),
    ("เปิดเว็บไม่ต้องลง Software",
     "เข้าจากที่ไหนก็ได้ มีอินเทอร์เน็ตและ Browser พอ"),
    ("ปลอดภัย รองรับ HTTPS",
     "Login + cookie signed + API auth — เหมาะกับการใช้ในองค์กร"),
    ("ปรับแต่งได้ตามต้องการ",
     "ตั้งชื่อตู้ ตั้งค่า threshold ตั้ง Telegram bot ของลูกค้าเองได้"),
]
for head, body in reasons:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(head + " — ")
    r.font.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = BRAND_GREEN
    r2 = p.add_run(body)
    r2.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════
# 9. ภาพรวมแพ็กเกจ
# ════════════════════════════════════════════════════════════════════════

add_h1("9. สิ่งที่ลูกค้าจะได้รับ")

add_feature_box("ติดตั้งครั้งแรก (Initial Setup)", [
    "ติดตั้งระบบบน Server ของลูกค้า หรือใช้ Server กลางของ EDS",
    "Setup MQTT broker + เชื่อมต่อกับตู้แต่ละตัว",
    "ตั้งค่า Domain + HTTPS",
    "สร้าง Username/Password ให้ทีมงาน",
    "เทรนนิ่งการใช้งาน (1-2 ชั่วโมง)",
])

add_feature_box("การดูแลรักษา (Ongoing Support)", [
    "Update ระบบเมื่อมีฟีเจอร์ใหม่",
    "ดูแลฐานข้อมูล + Backup รายสัปดาห์",
    "Technical Support ผ่าน Email / Line / Telegram",
    "Bug fix ฟรีในรอบ 12 เดือนแรก",
])

add_feature_box("คุณค่าเพิ่ม", [
    "ลด downtime — เห็นปัญหาเร็วขึ้น ตอบสนองได้ใน 5 นาที (จากเดิมอาจ 1-2 ชม.)",
    "ลดต้นทุนการตรวจสอบหน้างาน — ไม่ต้องส่งคนไปดูทุกครั้ง",
    "เพิ่มความน่าเชื่อถือต่อลูกค้าปลายทาง — ตู้พร้อมใช้งานมากขึ้น",
    "ข้อมูลรายงานพร้อมเสนอผู้บริหาร / นักลงทุน",
])

# ════════════════════════════════════════════════════════════════════════
# Footer / Contact
# ════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ติดต่อสอบถาม / นัดเดโมระบบ")
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = BRAND_GREEN

add_title("EDS Engineering Solutions", size=13, color=DARK, bold=True, space_after=2)
add_title("FlexxFast EV Monitoring Platform", size=11, color=MUTED, bold=False)

# Save
doc.save(OUT)
print(f"[OK] Saved: {OUT}")
print(f"     Size:  {OUT.stat().st_size / 1024:.1f} KB")
